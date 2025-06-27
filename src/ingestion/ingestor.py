from typing import List, Dict
from pathlib import Path
import io
import fitz  # PyMuPDF
import docx
import hashlib
import datetime

from src.config import settings

if settings.USE_ONEDRIVE:
    try:
        from .onedrive_client import OneDriveClient
    except Exception:
        OneDriveClient = None

SUPPORTED_EXTENSIONS = [".pdf", ".docx", ".txt"]

def extract_text_from_pdf(file_path: Path) -> str:
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def extract_text_from_docx(file_path: Path) -> str:
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_txt(file_path: Path) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

def extract_text(file_path: Path) -> str:
    if file_path.suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif file_path.suffix == ".docx":
        return extract_text_from_docx(file_path)
    elif file_path.suffix == ".txt":
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_path.suffix}")

def extract_text_from_bytes(filename: str, data: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        doc = fitz.open(stream=data, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    elif ext == ".docx":
        doc = docx.Document(io.BytesIO(data))
        return "\n".join([para.text for para in doc.paragraphs])
    elif ext == ".txt":
        return data.decode("utf-8")
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def compute_hash(data: bytes) -> str:
    """Devuelve un hash SHA1 del contenido para identificarlo de forma estable."""
    return hashlib.sha1(data).hexdigest()

def process_documents(input_folder: Path, output_folder: Path, save_to_disk: bool = True) -> List[Dict]:
    """Procesa documentos locales y, opcionalmente, archivos remotos de OneDrive.

    Si ``save_to_disk`` es ``False`` los textos extraídos no se guardan en
    ``output_folder`` y solo se devuelven en memoria.
    """
    if save_to_disk:
        output_folder.mkdir(parents=True, exist_ok=True)
    input_folder.mkdir(parents=True, exist_ok=True)

    processed_docs = []

    if settings.USE_ONEDRIVE and OneDriveClient:
        try:
            client = OneDriveClient(
                settings.ONEDRIVE_CLIENT_ID,
                settings.ONEDRIVE_CLIENT_SECRET,
                settings.ONEDRIVE_TENANT_ID,
            )
            for name, item_id, modified, data in client.iter_files(
                drive_id=settings.ONEDRIVE_DRIVE_ID,
                folder_path=settings.ONEDRIVE_FOLDER,
            ):
                if Path(name).suffix.lower() not in SUPPORTED_EXTENSIONS:
                    continue
                print(f"Procesando remoto: {name}")
                content = extract_text_from_bytes(name, data)
                file_hash = compute_hash(data)
                metadata = {
                    "doc_id": file_hash,
                    "filename": name,
                    "created": modified or datetime.datetime.now().isoformat(),
                    "source": f"OneDrive:{settings.ONEDRIVE_FOLDER}/{name}",
                    "onedrive_id": item_id,
                }
                if save_to_disk:
                    with open(output_folder / f"{file_hash}.txt", "w", encoding="utf-8") as f_out:
                        f_out.write(content)
                processed_docs.append({"text": content, "metadata": metadata})
        except Exception as e:
            print(f"Error sincronizando OneDrive: {e}")

    for file in input_folder.iterdir():
        if file.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue

        print(f"Procesando: {file.name}")
        try:
            content = extract_text(file)
            data_bytes = content.encode("utf-8")
            file_hash = compute_hash(data_bytes)
            metadata = {
                "doc_id": file_hash,
                "filename": file.name,
                "created": datetime.datetime.now().isoformat(),
                "source": str(file.resolve())
            }


            if save_to_disk:
                with open(output_folder / f"{file_hash}.txt", "w", encoding="utf-8") as f_out:
                    f_out.write(content)

            processed_docs.append({"text": content, "metadata": metadata})

        except Exception as e:
            print(f"Error procesando {file.name}: {e}")
    
    return processed_docs
